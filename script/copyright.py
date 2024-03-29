import pathlib as p
import re
import typing as t

from docx import Document
from docx.shared import Pt, RGBColor
from pygments.lexers import guess_lexer_for_filename
from pygments.styles import get_style_by_name

if t.TYPE_CHECKING:
    import typing_extensions as te

    P = te.ParamSpec('P')


Path = t.Union[str, p.Path]


class Word:
    '''
    Reference:
        - http://www.gov.cn/zhengce/2020-12/26/content_5574414.htm
    '''

    def __init__(
        self,
        root: Path = '.', style: str = 'default', page_break: bool = True,
        font_name: str = 'Times New Roman', font_size: int = 9,
    ) -> None:
        self._root = p.Path(root).absolute()
        self._default = {
            'style': style,
            'page_break': page_break,
        }
        self._doc = Document()
        font = self._doc.styles['Normal'].font
        font.name = font_name
        font.size = Pt(font_size)

    @classmethod
    def fromConfig(cls, *args: 'P.args', **kwargs: 'P.kwargs') -> 'te.Self':
        return cls(*args, **kwargs)

    @property
    def root(self) -> p.Path:
        return self._root

    def add(
        self,
        path: Path, plain: bool = False,
        style: t.Optional[str] = None, page_break: t.Optional[bool] = None,
        title: t.Optional[str] = None,
    ) -> 'te.Self':
        if plain:
            return self._add_plain(path, page_break, title)
        path = p.Path(path).absolute()
        code = self._strip(path.read_text())
        lexer = guess_lexer_for_filename(path.name, code)
        styles = dict(get_style_by_name(style or self._default['style']))
        # heading
        self._doc.add_heading(title or path.relative_to(self._root).as_posix(), 1)
        # paragraph
        paragraph = self._doc.add_paragraph()
        for type, value in lexer.get_tokens(code):
            style = styles.get(type, {})
            run = paragraph.add_run(value)
            # bold, italic, underline
            run.bold = style.get('bold', False)
            run.italic = style.get('italic', False)
            run.underline = style.get('underline', False)
            # color
            color = style.get('color', None)
            if color is not None:
                run.font.color.rgb = RGBColor.from_string(color)
        # page break
        if page_break or self._default['page_break']:
            self._doc.add_page_break()
        return self

    def save(self, path: Path) -> 'te.Self':
        self._doc.save(path)
        return self

    def _add_plain(
        self,
        path: Path,
        page_break: t.Optional[bool] = None, title: t.Optional[str] = None,
    ) -> 'te.Self':
        path = p.Path(path).absolute()
        self._doc.add_heading(title or path.relative_to(self._root).as_posix(), 1)
        self._doc.add_paragraph(self._strip(path.read_text()))
        if page_break or self._default['page_break']:
            self._doc.add_page_break()
        return self

    def _strip(self, string: str) -> str:
        return re.sub(r'\n+', '\n', string.strip())


if __name__ == '__main__':
    is_valid: t.Callable[[p.Path], bool] \
        = lambda path: 'type: [embed, 7z]' not in path.read_text()

    word = Word.fromConfig(root='.', style='friendly', page_break=False, font_size=7)
    word.add(word.root/'pyproject.toml')
    word.add(word.root/'foam.py')
    for path in (word.root/'foam').rglob('*.py'):
        word.add(path)
    for path in (word.root/'script').rglob('*.py'):
        word.add(path)
    for path in (word.root/'extra/tutorial/tutorials').rglob('*.yaml'):
        if is_valid(path):
            word.add(path)
    for path in (word.root/'extra/tutorial/third_party').rglob('*'):
        if path.is_file() and path.suffix in {'.C', '.H', '.yaml', '.sh'}:
            if is_valid(path):
                word.add(path)
    word.save('copyright.docx')
